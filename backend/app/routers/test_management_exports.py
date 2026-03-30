"""Export endpoints for test management — Excel, Word, PDF."""
from __future__ import annotations

import io
import uuid
from typing import Optional

import structlog
from fastapi import APIRouter, Depends, Query
from fastapi.responses import Response, StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.deps import get_current_active_user
from app.db.postgres import get_db
from app.models.postgres import ManagedTestCase, TestPlan, TestPlanItem, TestStrategy, User

logger = structlog.get_logger(__name__)

router = APIRouter()


# ── Excel export: test cases ──────────────────────────────────────────────────

@router.get("/cases/export/excel")
async def export_test_cases_excel(
    project_id: Optional[uuid.UUID] = None,
    status: Optional[str] = None,
    test_type: Optional[str] = None,
    priority: Optional[str] = None,
    search: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export test cases to an Excel (.xlsx) file."""
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    # Explicitly select only the columns we need — avoids breaking if a new
    # column (e.g. suite_name from migration 0013) hasn't been applied yet.
    stmt = select(
        ManagedTestCase.title,
        ManagedTestCase.status,
        ManagedTestCase.test_type,
        ManagedTestCase.priority,
        ManagedTestCase.severity,
        ManagedTestCase.feature_area,
        ManagedTestCase.objective,
        ManagedTestCase.preconditions,
        ManagedTestCase.expected_result,
        ManagedTestCase.is_automated,
        ManagedTestCase.automation_status,
        ManagedTestCase.ai_generated,
        ManagedTestCase.ai_quality_score,
        ManagedTestCase.last_execution_status,
        ManagedTestCase.tags,
        ManagedTestCase.steps,
        ManagedTestCase.version,
        ManagedTestCase.created_at,
    ).order_by(ManagedTestCase.created_at.desc())
    if project_id:
        stmt = stmt.where(ManagedTestCase.project_id == project_id)
    if status:
        stmt = stmt.where(ManagedTestCase.status == status)
    if test_type:
        stmt = stmt.where(ManagedTestCase.test_type == test_type)
    if priority:
        stmt = stmt.where(ManagedTestCase.priority == priority)
    if search:
        stmt = stmt.where(ManagedTestCase.title.ilike(f"%{search}%"))

    cases = (await db.execute(stmt)).all()
    logger.info("exporting_test_cases_excel", count=len(cases), project_id=str(project_id) if project_id else None)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Test Cases"

    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)

    headers = [
        "Title", "Status", "Type", "Priority", "Severity",
        "Feature Area", "Objective", "Preconditions", "Steps", "Expected Result",
        "Automated", "Automation Status", "AI Generated", "AI Quality Score",
        "Last Execution Status", "Tags", "Version", "Created At",
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    ws.row_dimensions[1].height = 20

    for row_idx, tc in enumerate(cases, start=2):
        steps_text = ""
        if tc.steps:
            steps_text = "; ".join(
                f"Step {s.get('step_number', i+1)}: {s.get('action', '')}"
                for i, s in enumerate(tc.steps)
            )
        tags_text = ", ".join(tc.tags) if tc.tags else ""

        ws.cell(row=row_idx, column=1, value=tc.title)
        ws.cell(row=row_idx, column=2, value=tc.status)
        ws.cell(row=row_idx, column=3, value=tc.test_type)
        ws.cell(row=row_idx, column=4, value=tc.priority)
        ws.cell(row=row_idx, column=5, value=tc.severity)
        ws.cell(row=row_idx, column=6, value=tc.feature_area or "")
        ws.cell(row=row_idx, column=7, value=tc.objective or "")
        ws.cell(row=row_idx, column=8, value=tc.preconditions or "")
        ws.cell(row=row_idx, column=9, value=steps_text)
        ws.cell(row=row_idx, column=10, value=tc.expected_result or "")
        ws.cell(row=row_idx, column=11, value="Yes" if tc.is_automated else "No")
        ws.cell(row=row_idx, column=12, value=tc.automation_status)
        ws.cell(row=row_idx, column=13, value="Yes" if tc.ai_generated else "No")
        ws.cell(row=row_idx, column=14, value=tc.ai_quality_score)
        ws.cell(row=row_idx, column=15, value=tc.last_execution_status or "")
        ws.cell(row=row_idx, column=16, value=tags_text)
        ws.cell(row=row_idx, column=17, value=tc.version)
        ws.cell(row=row_idx, column=18, value=tc.created_at.strftime("%Y-%m-%d") if tc.created_at else "")

    # Auto-fit column widths
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    buf = io.BytesIO()
    wb.save(buf)

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=test-cases.xlsx"},
    )


# ── Word/PDF export helpers ───────────────────────────────────────────────────

def _safe(value) -> str:
    if value is None:
        return ""
    return str(value)


def _list_to_str(items) -> str:
    if not items:
        return ""
    if isinstance(items, list):
        return "\n".join(f"• {str(i)}" for i in items)
    return str(items)


# ── Word export: test plan ────────────────────────────────────────────────────

@router.get("/plans/{plan_id}/export/word")
async def export_test_plan_word(
    plan_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export a test plan to a Word (.docx) document."""
    from docx import Document
    from docx.shared import RGBColor

    plan = (await db.execute(select(TestPlan).where(TestPlan.id == plan_id))).scalar_one_or_none()
    if not plan:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Test plan not found")

    items_result = await db.execute(
        select(TestPlanItem, ManagedTestCase)
        .join(ManagedTestCase, TestPlanItem.test_case_id == ManagedTestCase.id)
        .where(TestPlanItem.plan_id == plan_id)
        .order_by(TestPlanItem.order_index)
    )
    items = items_result.all()

    logger.info("exporting_plan_word", plan_id=str(plan_id), item_count=len(items))

    doc = Document()

    # Title
    title = doc.add_heading(_safe(plan.name), level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph(f"Status: {_safe(plan.status)}")
    if plan.description:
        doc.add_paragraph(_safe(plan.description))
    if plan.objective:
        doc.add_heading("Objective", level=2)
        doc.add_paragraph(_safe(plan.objective))

    # Schedule
    if plan.planned_start_date or plan.planned_end_date:
        doc.add_heading("Schedule", level=2)
        sched_para = doc.add_paragraph()
        if plan.planned_start_date:
            sched_para.add_run(f"Planned Start: {plan.planned_start_date.strftime('%Y-%m-%d')}  ")
        if plan.planned_end_date:
            sched_para.add_run(f"Planned End: {plan.planned_end_date.strftime('%Y-%m-%d')}")

    # Summary
    doc.add_heading("Summary", level=2)
    summary = doc.add_paragraph()
    summary.add_run(f"Total Cases: {plan.total_cases}   ")
    summary.add_run(f"Executed: {plan.executed_cases}   ")
    summary.add_run(f"Passed: {plan.passed_cases}   ")
    summary.add_run(f"Failed: {plan.failed_cases}   ")
    summary.add_run(f"Blocked: {plan.blocked_cases}")

    # Test cases table
    if items:
        doc.add_heading("Test Cases", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Shading"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["#", "Title", "Priority", "Type", "Status", "Execution Status"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True

        for idx, (item, tc) in enumerate(items, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = _safe(tc.title)
            row_cells[2].text = _safe(item.priority_override or tc.priority)
            row_cells[3].text = _safe(tc.test_type)
            row_cells[4].text = _safe(tc.status)
            row_cells[5].text = _safe(item.execution_status)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=test-plan-{plan_id}.docx"},
    )


# ── PDF export: test plan ─────────────────────────────────────────────────────

@router.get("/plans/{plan_id}/export/pdf")
async def export_test_plan_pdf(
    plan_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export a test plan to a PDF document."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    plan = (await db.execute(select(TestPlan).where(TestPlan.id == plan_id))).scalar_one_or_none()
    if not plan:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Test plan not found")

    items_result = await db.execute(
        select(TestPlanItem, ManagedTestCase)
        .join(ManagedTestCase, TestPlanItem.test_case_id == ManagedTestCase.id)
        .where(TestPlanItem.plan_id == plan_id)
        .order_by(TestPlanItem.order_index)
    )
    items = items_result.all()

    logger.info("exporting_plan_pdf", plan_id=str(plan_id), item_count=len(items))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle("Heading1", parent=styles["Heading1"], textColor=colors.HexColor("#1E40AF"), spaceAfter=6)
    heading2_style = ParagraphStyle("Heading2", parent=styles["Heading2"], textColor=colors.HexColor("#1E3A5F"), spaceAfter=4)
    body_style = styles["Normal"]

    story = [
        Paragraph(_safe(plan.name), heading_style),
        Paragraph(f"Status: {_safe(plan.status)} | Total Cases: {plan.total_cases}", body_style),
        Spacer(1, 6*mm),
    ]

    if plan.description:
        story += [Paragraph("Description", heading2_style), Paragraph(_safe(plan.description), body_style), Spacer(1, 4*mm)]

    if plan.objective:
        story += [Paragraph("Objective", heading2_style), Paragraph(_safe(plan.objective), body_style), Spacer(1, 4*mm)]

    # Execution summary
    story += [
        Paragraph("Execution Summary", heading2_style),
        Paragraph(
            f"Executed: {plan.executed_cases} / {plan.total_cases} | "
            f"Passed: {plan.passed_cases} | Failed: {plan.failed_cases} | Blocked: {plan.blocked_cases}",
            body_style,
        ),
        Spacer(1, 6*mm),
    ]

    if items:
        story.append(Paragraph("Test Cases", heading2_style))
        table_data = [["#", "Title", "Priority", "Type", "Exec Status"]]
        for idx, (item, tc) in enumerate(items, start=1):
            table_data.append([
                str(idx),
                (_safe(tc.title)[:60] + "…") if len(_safe(tc.title)) > 60 else _safe(tc.title),
                _safe(item.priority_override or tc.priority),
                _safe(tc.test_type),
                _safe(item.execution_status),
            ])

        tbl = Table(table_data, colWidths=[10*mm, 85*mm, 25*mm, 25*mm, 30*mm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(tbl)

    doc.build(story)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=test-plan-{plan_id}.pdf"},
    )


# ── Word export: test strategy ────────────────────────────────────────────────

@router.get("/strategies/{strategy_id}/export/word")
async def export_test_strategy_word(
    strategy_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export a test strategy to a Word (.docx) document."""
    from docx import Document
    from docx.shared import RGBColor

    strategy = (await db.execute(select(TestStrategy).where(TestStrategy.id == strategy_id))).scalar_one_or_none()
    if not strategy:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Test strategy not found")

    logger.info("exporting_strategy_word", strategy_id=str(strategy_id))

    doc = Document()
    title = doc.add_heading(_safe(strategy.name), level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    doc.add_paragraph(f"Version: {_safe(strategy.version_label)} | Status: {_safe(strategy.status)}")

    sections = [
        ("Objective", strategy.objective),
        ("Scope", strategy.scope),
        ("Out of Scope", strategy.out_of_scope),
        ("Test Approach", strategy.test_approach),
        ("Automation Approach", strategy.automation_approach),
        ("Defect Management", strategy.defect_management),
    ]
    for heading, content in sections:
        if content:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(_safe(content))

    if strategy.entry_criteria:
        doc.add_heading("Entry Criteria", level=2)
        for criterion in strategy.entry_criteria:
            doc.add_paragraph(f"• {criterion}", style="List Bullet")

    if strategy.exit_criteria:
        doc.add_heading("Exit Criteria", level=2)
        for criterion in strategy.exit_criteria:
            doc.add_paragraph(f"• {criterion}", style="List Bullet")

    if strategy.test_types:
        doc.add_heading("Test Types", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Type", "Priority", "Tools", "Coverage Target"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
        for tt in strategy.test_types:
            row_cells = table.add_row().cells
            row_cells[0].text = _safe(tt.get("type"))
            row_cells[1].text = _safe(tt.get("priority"))
            row_cells[2].text = _safe(tt.get("tools"))
            row_cells[3].text = f"{tt.get('coverage_target_pct', '')}%"

    if strategy.risk_assessment:
        doc.add_heading("Risk Assessment", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Risk", "Likelihood", "Impact", "Mitigation"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
        for risk in strategy.risk_assessment:
            row_cells = table.add_row().cells
            row_cells[0].text = _safe(risk.get("risk"))
            row_cells[1].text = _safe(risk.get("likelihood"))
            row_cells[2].text = _safe(risk.get("impact"))
            row_cells[3].text = _safe(risk.get("mitigation"))

    if strategy.environments:
        doc.add_heading("Environments", level=2)
        for env in strategy.environments:
            doc.add_paragraph(f"• {env.get('name', '')} ({env.get('type', '')}) — {env.get('purpose', '')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=test-strategy-{strategy_id}.docx"},
    )


# ── PDF export: test strategy ─────────────────────────────────────────────────

@router.get("/strategies/{strategy_id}/export/pdf")
async def export_test_strategy_pdf(
    strategy_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export a test strategy to a PDF document."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    strategy = (await db.execute(select(TestStrategy).where(TestStrategy.id == strategy_id))).scalar_one_or_none()
    if not strategy:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Test strategy not found")

    logger.info("exporting_strategy_pdf", strategy_id=str(strategy_id))

    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buf, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    heading1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=colors.HexColor("#1E40AF"), spaceAfter=6)
    heading2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor("#1E3A5F"), spaceAfter=4)
    body = styles["Normal"]

    story = [
        Paragraph(_safe(strategy.name), heading1),
        Paragraph(f"Version: {_safe(strategy.version_label)} | Status: {_safe(strategy.status)}", body),
        Spacer(1, 6*mm),
    ]

    text_sections = [
        ("Objective", strategy.objective),
        ("Scope", strategy.scope),
        ("Out of Scope", strategy.out_of_scope),
        ("Test Approach", strategy.test_approach),
        ("Automation Approach", strategy.automation_approach),
        ("Defect Management", strategy.defect_management),
    ]
    for heading, content in text_sections:
        if content:
            story += [Paragraph(heading, heading2), Paragraph(_safe(content), body), Spacer(1, 4*mm)]

    if strategy.entry_criteria:
        story.append(Paragraph("Entry Criteria", heading2))
        for c in strategy.entry_criteria:
            story.append(Paragraph(f"• {c}", body))
        story.append(Spacer(1, 4*mm))

    if strategy.exit_criteria:
        story.append(Paragraph("Exit Criteria", heading2))
        for c in strategy.exit_criteria:
            story.append(Paragraph(f"• {c}", body))
        story.append(Spacer(1, 4*mm))

    if strategy.risk_assessment:
        story.append(Paragraph("Risk Assessment", heading2))
        risk_data = [["Risk", "Likelihood", "Impact", "Mitigation"]]
        for risk in strategy.risk_assessment:
            risk_data.append([
                _safe(risk.get("risk")),
                _safe(risk.get("likelihood")),
                _safe(risk.get("impact")),
                _safe(risk.get("mitigation")),
            ])
        tbl = Table(risk_data, colWidths=[50*mm, 25*mm, 25*mm, 60*mm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story += [tbl, Spacer(1, 4*mm)]

    doc_pdf.build(story)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=test-strategy-{strategy_id}.pdf"},
    )


# ── Test Suites: list suites and their test cases ─────────────────────────────

@router.get("/suites")
async def list_test_suites(
    project_id: Optional[uuid.UUID] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Return test suites grouped by suite_name, combining automation test_cases
    (from ingested runs) and manually authored managed_test_cases.
    """
    from sqlalchemy import text as sa_text

    # Automation test cases
    auto_where = "AND tr.project_id = :project_id" if project_id else ""
    auto_params: dict = {"project_id": project_id} if project_id else {}
    auto_query = sa_text(f"""
        SELECT
            tc.suite_name,
            COUNT(*) AS test_count,
            COUNT(*) FILTER (WHERE tc.status = 'PASSED') AS passed_count,
            COUNT(*) FILTER (WHERE tc.status = 'FAILED') AS failed_count,
            MAX(tr.created_at) AS last_run_at
        FROM test_cases tc
        JOIN test_runs tr ON tc.test_run_id = tr.id
        WHERE tc.suite_name IS NOT NULL AND tc.suite_name != ''
          {auto_where}
        GROUP BY tc.suite_name
    """)
    auto_rows = (await db.execute(auto_query, auto_params)).fetchall()

    # Manual managed test cases (suite_name added in migration 0013)
    manual_where = "AND project_id = :project_id" if project_id else ""
    manual_query = sa_text(f"""
        SELECT
            suite_name,
            COUNT(*) AS test_count,
            0 AS passed_count,
            0 AS failed_count,
            MAX(created_at) AS last_run_at
        FROM managed_test_cases
        WHERE suite_name IS NOT NULL AND suite_name != ''
          {manual_where}
        GROUP BY suite_name
    """)
    try:
        manual_rows = (await db.execute(manual_query, auto_params)).fetchall()
    except Exception as exc:
        logger.warning("managed_test_cases.suite_name not available, skipping manual suites", error=str(exc))
        await db.rollback()
        manual_rows = []

    # Merge both sources by suite_name
    merged: dict[str, dict] = {}
    for row in auto_rows:
        merged[row.suite_name] = {
            "suite_name": row.suite_name,
            "test_count": row.test_count,
            "passed_count": row.passed_count,
            "failed_count": row.failed_count,
            "last_run_at": row.last_run_at,
        }
    for row in manual_rows:
        if row.suite_name in merged:
            merged[row.suite_name]["test_count"] += row.test_count
            if row.last_run_at and (
                merged[row.suite_name]["last_run_at"] is None
                or row.last_run_at > merged[row.suite_name]["last_run_at"]
            ):
                merged[row.suite_name]["last_run_at"] = row.last_run_at
        else:
            merged[row.suite_name] = {
                "suite_name": row.suite_name,
                "test_count": row.test_count,
                "passed_count": row.passed_count,
                "failed_count": row.failed_count,
                "last_run_at": row.last_run_at,
            }

    result = sorted(merged.values(), key=lambda x: x["test_count"], reverse=True)
    logger.info("listing_test_suites", count=len(result), project_id=str(project_id) if project_id else None)

    return [
        {
            "suite_name": s["suite_name"],
            "test_count": s["test_count"],
            "passed_count": s["passed_count"],
            "failed_count": s["failed_count"],
            "last_run_at": s["last_run_at"].isoformat() if s["last_run_at"] else None,
            "pass_rate": round(s["passed_count"] / s["test_count"] * 100, 1) if s["test_count"] > 0 else None,
        }
        for s in result
    ]


@router.get("/suites/{suite_name}/cases")
async def get_suite_test_cases(
    suite_name: str,
    project_id: Optional[uuid.UUID] = Query(None),
    limit: int = Query(100, ge=1, le=500),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Return test cases for a given suite_name from both automation runs and managed test cases."""
    from app.models.postgres import TestCase, TestRun

    # Automation test cases
    auto_stmt = (
        select(TestCase)
        .join(TestRun, TestCase.test_run_id == TestRun.id)
        .where(TestCase.suite_name == suite_name)
        .order_by(TestCase.created_at.desc())
        .limit(limit)
    )
    if project_id:
        auto_stmt = auto_stmt.where(TestRun.project_id == project_id)

    auto_cases = (await db.execute(auto_stmt)).scalars().all()

    # Manual managed test cases
    manual_stmt = (
        select(ManagedTestCase)
        .where(ManagedTestCase.suite_name == suite_name)
        .order_by(ManagedTestCase.created_at.desc())
        .limit(limit)
    )
    if project_id:
        manual_stmt = manual_stmt.where(ManagedTestCase.project_id == project_id)

    manual_cases = (await db.execute(manual_stmt)).scalars().all()

    result = []
    for tc in auto_cases:
        result.append({
            "id": str(tc.id),
            "test_name": tc.test_name,
            "suite_name": tc.suite_name,
            "status": tc.status,
            "duration_ms": tc.duration_ms,
            "class_name": tc.class_name,
            "package_name": tc.package_name,
            "created_at": tc.created_at.isoformat() if tc.created_at else None,
            "source": "automation",
        })
    for tc in manual_cases:
        result.append({
            "id": str(tc.id),
            "test_name": tc.title,
            "suite_name": tc.suite_name,
            "status": tc.last_execution_status or tc.status,
            "duration_ms": None,
            "class_name": tc.feature_area,
            "package_name": None,
            "created_at": tc.created_at.isoformat() if tc.created_at else None,
            "source": "manual",
        })

    # Sort combined by created_at descending
    result.sort(key=lambda x: x["created_at"] or "", reverse=True)
    return result[:limit]
